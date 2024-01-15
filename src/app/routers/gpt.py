from fastapi import FastAPI, Depends
import logging
from extract_the_result import load_action_prompt
from extract_the_result import chat_with_prompt_for_pic, chat_with_prompt_for_exam
from extract_the_result import AiCmdEnum
# import urllib.parse
import requests
import uvicorn
import json
import os
from typing import List
from fastapi import File, UploadFile
import config.constants as ct
import time
from app.utils.service_result import ServiceResult
from app.schemas.reqhistory import ReqItemCreate
from app.services.reqhistory import ReqHistoryService
from app.models.dao_reqhistory import RequestHistoryCRUD
from config.database import get_db
from docx import Document
import re



os.environ['http_proxy'] = ct.OPEN_AI_HTTP_PROXY
os.environ['https_proxy'] = ct.OPEN_AI_HTTP_PROXY
# 配置logging
logger = logging.getLogger()
logger.setLevel(logging.DEBUG)

app2 = FastAPI(
    title="OpenAI微服务",
    version="0.1.0",
)


def extract_string_between_identifiers(input_string, start_identifier, end_identifier):
    pattern = re.compile(f'{re.escape(start_identifier)}(.*?){re.escape(end_identifier)}', re.DOTALL)
    match = pattern.search(input_string)
    if match:
        return match.group(1)
    else:
        return None


def replace_placeholder(doc, placeholder, replacement):
    # 打开 Word 文档
    # doc = Document(doc_path)

    # 遍历文档中的段落
    for paragraph in doc.paragraphs:
        # 在段落文本中查找占位符
        if placeholder in paragraph.text:
            # 替换占位符为指定的内容
            paragraph.text = paragraph.text.replace(placeholder, replacement)

    # 遍历文档中的表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 在表格单元格中查找占位符
                if placeholder in cell.text:
                    # 替换占位符为指定的内容
                    cell.text = cell.text.replace(placeholder, replacement)

    # 保存修改后的文档
    # doc.save("modified_document.docx")


@app2.put("/ai/tutor")
def ai_tutor(course_objs: str):
    res = {}
    try:
        exam_prompt = load_action_prompt(AiCmdEnum.exam)
        # ct.OPEN_AI_MODLE should be gpt-3.5-turbo-1106 or gpt-3.5-turbo or
        resp = chat_with_prompt_for_exam(exam_prompt, mode=ct.OPEN_AI_MODEL, temperature=0.0, content=course_objs)
        # resp_obj = json.loads(resp)
        # content = resp_obj.get('content')
        result = {
            "code": 200,
            "content": resp
        }

        # 打开 Word 文档
        doc_template_path = 'E:\\桌面\\exam_paper\\tpl.docx'
        target_path = 'E:\\桌面\\exam_paper\\gen.docx'

        # 创建文档
        doc = Document(doc_template_path)
        items = ''
        sections_str = result.get('content')
        sections_str = extract_string_between_identifiers(sections_str, '```json\n', '\n```')
        sections = json.loads(sections_str)
        # 填空题
        fill_in_blank = sections.get('fill-in-blank')
        for ind, ques in enumerate(fill_in_blank):
            q_it = f'{ind+1}.' + ques.get('question') + '\n'
            items = items + q_it
        place_holder = '[fill_blank]'
        replace_placeholder(doc, place_holder, items)
        # 选择题
        items = ''
        m_c = sections.get('multiple-choice')
        for ind, ques in enumerate(m_c):
            q_it = f'{ind+1}.' + ques.get('question') + '\n'
            opts = ques.get('options')
            choices = ''
            for idx, opt in enumerate(opts):
                opts_str = ''
                opts_str = chr(65 + idx) + '.' + opts_str + opt + '\n'
                choices = choices + opts_str

            items = items + q_it + choices
        place_holder = '[multiple_choice]'
        replace_placeholder(doc, place_holder, items)

        doc.save(target_path)

        res['code'] = 200
        res['content'] = result
    except Exception as err:
        logger.error(err)
        res = {'code': 500, 'content': f'{err}'}
    finally:
        return res


@app2.put("/gpt/")
async def gpt_llm(repid: str, chara: str, db: get_db = Depends(), memo='', model=''):
    logger.info(f'成功接收chara！')
    # 将chara转回字典
    chara = json.loads(chara)
    final = []
    for i in chara.keys():
        final.append(chara[i])
    logger.info(final)

    logger.info(f'openAI start!')
    ocr_prompt = load_action_prompt(AiCmdEnum.ocr)

    try:
        result = chat_with_prompt_for_pic(ocr_prompt, mode=ct.OPEN_AI_MODEL, temperature=0.0, content=final)

    except Exception as err:
        logger.error(err)
        result = 'OpenAI调用出现错误'

    logger.info(f'openAI ended')

    # 将结果保存入数据库
    logger.info(f'save db start!')
    dev_type = model
    external_data = {
        'model': dev_type,
        'status': ct.REQ_STATUS_PENDING,
        'result': ct.REQ_STATUS_PENDING,
        'requestts': int(time.time() * 1000),
        'memo': memo
    }
    item = ReqItemCreate(**external_data)

    outline_item = RequestHistoryCRUD(db).create_record(item)
    res = ServiceResult(outline_item)

    reqs = ReqHistoryService(db)
    reqs.update_item(res.value.id, result)
    logger.info(f'save db ended!')
    return result


@app2.post("/test/")
async def test(files: List[UploadFile] = File(...)):
    file_bytes = None
    for file in files:
        file_bytes = await file.read()
    img_files = [("files", file_bytes)]

    url = f"https://127.0.0.1:29083/ocr/"
    response = requests.post(url, files=img_files, verify=False)

    if response.status_code == requests.codes.ok:
        print("PUT请求成功！")
    else:
        print("PUT请求失败！")

    li_data = response.json()
    data = json.loads(li_data)
    return data


if __name__ == '__main__':
    # 启动fastAPI
    logging.info(f'*********************  GPT AI services  ********************')
    uvicorn.run(app2,
                host='0.0.0.0',
                port=29082,
                ssl_keyfile=ct.SCHEDULE_KEY,
                ssl_certfile=ct.SCHEDULE_CER
                )
